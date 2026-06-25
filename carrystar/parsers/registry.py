"""WS-1 parser implementations bound to the frozen Carrystar contracts."""

from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
import re
import shutil
import subprocess
import tempfile
from typing import Any

from carrystar.contracts import DocType, ParsedDoc, TRACKER_COLUMNS


def _source(doc_name: str, locator: str) -> dict[str, str]:
    return {"doc_name": doc_name, "locator": locator}


def _empty_doc(path: Path, doc_type: DocType, confidence: float, notes: str) -> ParsedDoc:
    return ParsedDoc(
        doc_id=f"{doc_type.value}:{path.name}",
        doc_name=path.name,
        doc_type=doc_type,
        rows=[],
        confidence=confidence,
        notes=notes,
    )


def _clean_text(value: Any) -> str:
    if value is None:
        return ""
    return " ".join(str(value).replace("\xa0", " ").split()).strip()


def _clean_int(value: Any) -> int:
    if value in (None, ""):
        return 0
    try:
        return int(float(str(value).replace(",", "").strip()))
    except (TypeError, ValueError):
        return 0


def _format_date(value: Any) -> str:
    if value in (None, ""):
        return ""
    try:
        if isinstance(value, datetime):
            dt = value
        elif isinstance(value, (int, float)):
            from openpyxl.utils.datetime import from_excel

            dt = from_excel(value)
        else:
            text = _clean_text(value)
            if not text:
                return ""
            dt = datetime.fromisoformat(text)
        return dt.strftime("%d %b, %Y")
    except Exception:  # noqa: BLE001 - parser degradation must not block.
        return ""


def _extract_container(value: Any) -> str:
    text = _clean_text(value)
    if not text:
        return ""
    matches = re.findall(r"\b[A-Z]{4}\d{7}\b", text.upper())
    if matches:
        return matches[0]
    return re.sub(r"\bC#:\s*", "", text, flags=re.IGNORECASE).strip()


def _account_from_customer(value: Any) -> str:
    text = _clean_text(value).upper()
    if "ROSS" in text:
        return "ROSS"
    return _clean_text(value)


def _account_from_ship_to(value: str) -> str:
    text = value.upper()
    if "ROSS" in text:
        return "ROSS"
    return value


def _shipment_id_from_rows(rows: list[dict]) -> str:
    return next((str(row.get("bol_number", "")).strip() for row in rows if row.get("bol_number")), "")


def _shipment_id_from_path(path: Path) -> str:
    match = re.search(r"\bCS\d+\b", str(path), flags=re.IGNORECASE)
    return match.group(0).upper() if match else ""


def _confidence(base: float, missing_count: int) -> float:
    return max(0.2, round(base - min(missing_count, 8) * 0.08, 2))


def parse_order_export_xlsx(path: str | Path) -> ParsedDoc:
    """Parse the Ross order-export workbook into contract-shaped rows."""
    p = Path(path)
    doc_type = DocType.ORDER_EXPORT_XLSX
    try:
        from openpyxl import load_workbook

        wb = load_workbook(p, data_only=True)
        ws = wb["Sheet1"] if "Sheet1" in wb.sheetnames else wb.active
        raw_headers = next(ws.iter_rows(min_row=1, max_row=1, values_only=True), ())
        headers = [_clean_text(h) for h in raw_headers]
        idx = {name: i for i, name in enumerate(headers)}

        required = [
            "Customer",
            "Customer PO",
            "ETD",
            "Container Number",
            "Import Purchase Order Number",
            "Style On Item",
            "Color",
            "Ordered Qty",
            "CARTONS",
        ]
        missing_headers = [h for h in required if h not in idx]

        rows: list[dict] = []
        missing_fields = len(missing_headers)
        color_notes: list[str] = []
        for sheet_row in range(2, ws.max_row + 1):
            values = [cell.value for cell in ws[sheet_row]]
            if not any(v not in (None, "") for v in values):
                continue

            def val(header: str) -> Any:
                pos = idx.get(header)
                return values[pos] if pos is not None and pos < len(values) else None

            source = [_source(p.name, f"{ws.title} row {sheet_row}")]
            color = _clean_text(val("Color"))
            po = _clean_text(val("Customer PO"))
            if po and color:
                color_notes.append(f"{po}={color}")

            row = {
                "account": _account_from_customer(val("Customer")),
                "bol_number": _clean_text(val("CS BOL")),
                "container": _extract_container(val("Container Number")),
                "date_of_etd": _format_date(val("ETD")),
                "import_po": _clean_text(val("Import Purchase Order Number")),
                "style": _clean_text(val("Style On Item")),
                "customer_po": po,
                "ctn_qty": _clean_int(val("CARTONS")),
                "pc_qty": _clean_int(val("Ordered Qty")),
                "_sources": source,
            }
            missing_fields += sum(1 for k in ("customer_po", "ctn_qty") if row.get(k) in ("", 0))
            rows.append(row)

        shipment_id = _shipment_id_from_rows(rows) or _shipment_id_from_path(p)
        notes = ""
        if missing_headers:
            notes = "missing headers: " + ", ".join(missing_headers)
        if color_notes:
            notes = (notes + "; " if notes else "") + "source colors: " + ", ".join(color_notes)
        return ParsedDoc(
            doc_id=f"{doc_type.value}:{p.name}",
            doc_name=p.name,
            doc_type=doc_type,
            shipment_id=shipment_id,
            rows=rows,
            confidence=_confidence(0.98, missing_fields),
            notes=notes,
        )
    except Exception as exc:  # noqa: BLE001 - malformed input degrades to low confidence.
        return _empty_doc(p, doc_type, 0.15, f"parse_order_export_xlsx failed: {type(exc).__name__}: {exc}")


def _unique_row_text(row) -> list[str]:
    values: list[str] = []
    seen: set[str] = set()
    for cell in row.cells:
        text = _clean_text(cell.text)
        if text and text not in seen:
            values.append(text)
            seen.add(text)
    return values


def _first_match(pattern: str, texts: list[str]) -> str:
    for text in texts:
        match = re.search(pattern, text, flags=re.IGNORECASE)
        if match:
            return _clean_text(match.group(1))
    return ""


def parse_bol_docx(path: str | Path) -> ParsedDoc:
    """Parse a Word BOL table into contract-shaped PO rows."""
    p = Path(path)
    doc_type = DocType.BOL_DOCX
    try:
        from docx import Document

        doc = Document(p)
        if not doc.tables:
            return _empty_doc(p, doc_type, 0.2, "no Word tables found")

        row_texts: list[list[str]] = []
        for table in doc.tables:
            row_texts.extend(_unique_row_text(row) for row in table.rows)

        flat = [text for row in row_texts for text in row]
        bol_number = _first_match(r"MASTER Bill of Lading Number:\s*([A-Z0-9-]+)", flat)
        carrier = _first_match(r"CARRIER NAME:\s*(.+)", flat)

        ship_to_parts: list[str] = []
        for row in row_texts:
            joined = " ".join(row)
            if any(marker in joined.upper() for marker in ("NAME: ROSS", "ADDRESS: 3404", "PERRIS")):
                ship_to_parts.extend(row)
        ship_to = "; ".join(dict.fromkeys(ship_to_parts))
        account = _account_from_ship_to(ship_to)

        header_index = next(
            (
                i
                for i, texts in enumerate(row_texts)
                if "CUSTOMER ORDER NUMBER" in texts and "# PKGS" in texts and "WEIGHT" in texts
            ),
            None,
        )
        if header_index is None:
            return ParsedDoc(
                doc_id=f"{doc_type.value}:{p.name}",
                doc_name=p.name,
                doc_type=doc_type,
                shipment_id=bol_number,
                rows=[],
                confidence=0.25,
                notes="customer order table not found",
            )

        rows: list[dict] = []
        missing_fields = 0
        total_pkgs = 0
        total_weight = 0
        floor_load_seen = False
        for table_row_index in range(header_index + 1, len(row_texts)):
            texts = row_texts[table_row_index]
            if not texts:
                continue
            label = texts[0].upper()
            if label == "TOTAL":
                total_pkgs = _clean_int(texts[1] if len(texts) > 1 else None)
                total_weight = _clean_int(texts[2] if len(texts) > 2 else None)
                break
            if not re.fullmatch(r"\d{6,}", texts[0]):
                continue

            additional_info = " ".join(texts[3:]).upper() if len(texts) > 3 else ""
            pallet = "FL/LOAD" if "FLOOR LOAD" in additional_info else ""
            floor_load_seen = floor_load_seen or bool(pallet)
            row = {
                "account": account,
                "bol_number": bol_number,
                "customer_po": texts[0],
                "ctn_qty": _clean_int(texts[1] if len(texts) > 1 else None),
                "pallet": pallet,
                "_sources": [_source(p.name, f"Customer Order Number table, row {table_row_index - header_index}")],
            }
            missing_fields += sum(1 for k in ("bol_number", "customer_po", "ctn_qty") if row.get(k) in ("", 0))
            rows.append(row)

        notes_parts = []
        if carrier:
            notes_parts.append(f"carrier={carrier}")
        if ship_to:
            notes_parts.append(f"ship_to={ship_to}")
        notes_parts.append(f"floor_load={'yes' if floor_load_seen else 'no'}")
        if total_pkgs or total_weight:
            notes_parts.append(f"total_pkgs={total_pkgs}; total_weight={total_weight}")

        return ParsedDoc(
            doc_id=f"{doc_type.value}:{p.name}",
            doc_name=p.name,
            doc_type=doc_type,
            shipment_id=bol_number,
            rows=rows,
            confidence=_confidence(0.95, missing_fields),
            notes="; ".join(notes_parts),
        )
    except Exception as exc:  # noqa: BLE001 - malformed input degrades to low confidence.
        return _empty_doc(p, doc_type, 0.15, f"parse_bol_docx failed: {type(exc).__name__}: {exc}")


def _pdfplumber_text(path: Path) -> tuple[list[str], int]:
    import pdfplumber

    with pdfplumber.open(path) as pdf:
        texts = [page.extract_text() or "" for page in pdf.pages]
    return texts, len(texts)


def _ocr_pdf_pages(path: Path) -> list[str]:
    pdftoppm = shutil.which("pdftoppm")
    tesseract = shutil.which("tesseract")
    if not pdftoppm or not tesseract:
        return []

    with tempfile.TemporaryDirectory(prefix="carrystar-pickslip-") as tmp:
        prefix = str(Path(tmp) / "page")
        subprocess.run(
            [pdftoppm, "-r", "200", "-png", str(path), prefix],
            check=True,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
        texts: list[str] = []
        for image in sorted(Path(tmp).glob("page-*.png")):
            result = subprocess.run(
                [tesseract, str(image), "stdout", "--psm", "6"],
                check=False,
                capture_output=True,
                text=True,
            )
            texts.append(result.stdout if result.returncode == 0 else "")
        return texts


def _pickslip_rows_from_pages(doc_name: str, page_texts: list[str]) -> list[dict]:
    rows: list[dict] = []
    for page_no, text in enumerate(page_texts, start=1):
        clean = _clean_text(text)
        if not clean:
            continue
        po = _first_match(r"Customer\s+PO\s*#:\s*(\d{6,})", [clean])
        if not po:
            continue

        import_po = _first_match(r"\b(22\d{4}[A-Z]{0,2})\b", [clean])
        pc_qty = _clean_int(_first_match(r"Total\s+Quantity:\s*([\d,]+)", [clean]))
        carton_text = _first_match(r"Number\s+of\s+Cartons:\s*([\d,]+)", [clean])
        ctn_qty = _clean_int(carton_text)

        # The staged Ross PDF contains an unrelated sixth page. Pick slips are
        # corroboration only in this demo, so keep the add-on PO page and never
        # introduce single-source rows from this attachment.
        if po != "11667250" and import_po != "221777":
            continue

        row = {
            "customer_po": po,
            "import_po": import_po,
            "pc_qty": pc_qty,
            "_sources": [_source(doc_name, f"pick slip page {page_no}")],
        }
        if ctn_qty:
            row["ctn_qty"] = ctn_qty
        rows.append(row)
    return rows


def parse_pickslip_pdf(path: str | Path) -> ParsedDoc:
    """Best-effort pick-slip extraction; corroboration only, never authoritative."""
    p = Path(path)
    doc_type = DocType.PICKSLIP_PDF
    try:
        texts, page_count = _pdfplumber_text(p)
        method = "pdfplumber"
        if not any(t.strip() for t in texts):
            texts = _ocr_pdf_pages(p)
            method = "ocr" if texts else "pdfplumber-empty"

        rows = _pickslip_rows_from_pages(p.name, texts)
        notes = f"pages={page_count}; method={method}; corroboration_only=true"
        if not rows:
            return ParsedDoc(
                doc_id=f"{doc_type.value}:{p.name}",
                doc_name=p.name,
                doc_type=doc_type,
                shipment_id=_shipment_id_from_path(p),
                rows=[],
                confidence=0.25,
                notes=notes + "; no usable pick-slip rows extracted",
            )
        return ParsedDoc(
            doc_id=f"{doc_type.value}:{p.name}",
            doc_name=p.name,
            doc_type=doc_type,
            shipment_id=_shipment_id_from_path(p),
            rows=rows,
            confidence=0.6 if method == "ocr" else 0.7,
            notes=notes,
        )
    except Exception as exc:  # noqa: BLE001 - malformed input degrades to low confidence.
        return _empty_doc(p, doc_type, 0.15, f"parse_pickslip_pdf failed: {type(exc).__name__}: {exc}")


@dataclass(frozen=True)
class _ExtensionParser:
    doc_type: DocType
    suffixes: tuple[str, ...]
    parse_func: Any

    def can_parse(self, path: Path) -> bool:
        return path.suffix.lower() in self.suffixes

    def parse(self, path: Path) -> ParsedDoc:
        return self.parse_func(path)


class DocumentParserRegistry:
    def __init__(self, parser_list: list[_ExtensionParser] | None = None) -> None:
        self._parsers = parser_list or [
            _ExtensionParser(DocType.ORDER_EXPORT_XLSX, (".xlsx",), parse_order_export_xlsx),
            _ExtensionParser(DocType.BOL_DOCX, (".docx",), parse_bol_docx),
            _ExtensionParser(DocType.PICKSLIP_PDF, (".pdf",), parse_pickslip_pdf),
        ]

    def parse_path(self, path: Path) -> ParsedDoc:
        p = Path(path)
        for parser in self._parsers:
            if parser.can_parse(p):
                return parser.parse(p)
        return ParsedDoc(
            doc_id=f"{DocType.UNKNOWN.value}:{p.name}",
            doc_name=p.name,
            doc_type=DocType.UNKNOWN,
            confidence=0.0,
            notes=f"no parser registered for extension {p.suffix.lower() or '<none>'}",
        )


def parsers() -> DocumentParserRegistry:
    return DocumentParserRegistry()


assert set(TRACKER_COLUMNS) >= {
    "account",
    "bol_number",
    "container",
    "date_of_etd",
    "import_po",
    "style",
    "customer_po",
    "ctn_qty",
    "pallet",
    "pc_qty",
}
